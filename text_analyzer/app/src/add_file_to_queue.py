import pika
from docx import Document
import json
import base64

connection = pika.BlockingConnection(pika.ConnectionParameters('localhost'))
channel = connection.channel()

arguments = {
'x-queue-type': 'classic'
}

#channel.queue_declare(queue='file_queue', durable=True, arguments=arguments)

def getText(doc):
    fullText = [para.text for para in doc.paragraphs]
    return '\n'.join(fullText)

def publishDocxFile(user_id, file_path):
    file = open(file_path, 'rb')
    doc = Document(file)
    fullText = ''.join([para.text for para in doc.paragraphs])
    file_content = fullText.encode('utf-8')
    encoded_content = base64.b64encode(file_content)

    properties = pika.BasicProperties(
        headers={'id': user_id},
        content_encoding='UTF-8'
    )

    channel.basic_publish(
        exchange='',
        routing_key='file_queue',
        body=encoded_content,
        properties=properties
    )

    print(" [x] Sent .docx file for user {}.".format(user_id))

user_id = 3
file_path = 'app/test_data/Трансформеры.docx'
publishDocxFile(user_id, file_path)